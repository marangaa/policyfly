from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

class InsuranceQuoteTemplate:
    def __init__(self):
        # Initialize document
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Set up document styles for consistent formatting"""
        # Header style
        header_style = self.doc.styles.add_style('Header Style', WD_STYLE_TYPE.PARAGRAPH)
        header_style.font.size = Pt(16)
        header_style.font.bold = True
        
        # Subheader style
        subheader_style = self.doc.styles.add_style('Subheader Style', WD_STYLE_TYPE.PARAGRAPH)
        subheader_style.font.size = Pt(14)
        subheader_style.font.bold = True
        
        # Normal text style
        normal_style = self.doc.styles.add_style('Normal Style', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.size = Pt(11)

    def add_company_header(self, company_name, logo_path=None):
        """Add company header with optional logo"""
        if logo_path:
            self.doc.add_picture(logo_path, width=Inches(2))
        
        header = self.doc.add_paragraph(company_name, style='Header Style')
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()  # Add spacing

    def add_quote_info(self, quote_number, valid_until):
        """Add quote reference information"""
        quote_info = self.doc.add_paragraph(style='Normal Style')
        quote_info.add_run('Quote Reference: ').bold = True
        quote_info.add_run(quote_number)
        quote_info.add_run('\nDate Generated: ').bold = True
        quote_info.add_run(datetime.now().strftime('%Y-%m-%d'))
        quote_info.add_run('\nValid Until: ').bold = True
        quote_info.add_run(valid_until)
        self.doc.add_paragraph()

    def add_client_info(self, client_data):
        """Add client information section"""
        self.doc.add_paragraph('Client Information', style='Subheader Style')
        client_info = self.doc.add_paragraph(style='Normal Style')
        
        for key, value in client_data.items():
            client_info.add_run(f'{key}: ').bold = True
            client_info.add_run(f'{value}\n')
        self.doc.add_paragraph()

    def add_coverage_details(self, coverage_items):
        """Add insurance coverage details"""
        self.doc.add_paragraph('Coverage Details', style='Subheader Style')
        
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Coverage Type'
        header_cells[1].text = 'Coverage Amount'
        header_cells[2].text = 'Premium'

        for item in coverage_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item['type']
            row_cells[1].text = item['amount']
            row_cells[2].text = item['premium']
        
        self.doc.add_paragraph()

    def add_terms_and_conditions(self, terms):
        """Add terms and conditions section"""
        self.doc.add_paragraph('Terms and Conditions', style='Subheader Style')
        terms_para = self.doc.add_paragraph(style='Normal Style')
        
        for term in terms:
            terms_para.add_run(f'• {term}\n')
        self.doc.add_paragraph()

    def add_total_premium(self, total_premium, payment_terms):
        """Add total premium and payment information"""
        self.doc.add_paragraph('Premium Summary', style='Subheader Style')
        premium_info = self.doc.add_paragraph(style='Normal Style')
        premium_info.add_run('Total Premium: ').bold = True
        premium_info.add_run(f'${total_premium:,.2f}\n')
        premium_info.add_run('Payment Terms: ').bold = True
        premium_info.add_run(payment_terms)
        self.doc.add_paragraph()

    def add_footer(self, contact_info):
        """Add footer with contact information"""
        footer = self.doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = contact_info
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def save_document(self, filename):
        """Save the document"""
        self.doc.save(filename)

# Example usage
def create_sample_quote():
    template = InsuranceQuoteTemplate()
    
    # Add company information
    template.add_company_header("ABC Insurance Company")
    
    # Add quote information
    template.add_quote_info("QT-2024-001", "2024-12-31")
    
    # Add client information
    client_data = {
        "Name": "John Doe",
        "Address": "123 Main Street, Anytown, ST 12345",
        "Phone": "(555) 555-5555",
        "Email": "john.doe@email.com",
        "Policy Type": "Auto Insurance"
    }
    template.add_client_info(client_data)
    
    # Add coverage details
    coverage_items = [
        {"type": "Liability Coverage", "amount": "$300,000", "premium": "$800"},
        {"type": "Collision Coverage", "amount": "$50,000", "premium": "$400"},
        {"type": "Comprehensive", "amount": "$50,000", "premium": "$300"}
    ]
    template.add_coverage_details(coverage_items)
    
    # Add terms and conditions
    terms = [
        "Coverage begins upon receipt of first payment",
        "30-day notice required for cancellation",
        "Claims must be reported within 24 hours of incident",
        "Deductibles apply per incident"
    ]
    template.add_terms_and_conditions(terms)
    
    # Add premium summary
    template.add_total_premium(1500.00, "Monthly payments of $125.00")
    
    # Add footer
    template.add_footer("ABC Insurance Company | Phone: (800) 555-0123 | Email: support@abcinsurance.com")
    
    # Save the document
    template.save_document("insurance_quote.docx")

if __name__ == "__main__":
    create_sample_quote()