#!/usr/bin/env python3
"""
Insurance Template Generator
This script creates a Word document template for insurance policies with conditional formatting.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def ensure_output_directory(directory="output"):
    """Creates the output directory if it doesn't exist."""
    if not os.path.exists(directory):
        os.makedirs(directory)
        print(f"Created output directory: {directory}")

def create_insurance_template(company_name="Sample Insurance Co."):
    """Creates an insurance document template with conditional sections."""
    try:
        print(f"Creating template for {company_name}...")
        doc = Document()
        
        # Set up styles
        print("Setting up document styles...")
        styles = doc.styles
        
        # Create header style
        header_style = styles.add_style('CustomHeader', WD_STYLE_TYPE.PARAGRAPH)
        header_style.font.size = Pt(16)
        header_style.font.bold = True
        header_style.font.color.rgb = RGBColor(0, 51, 102)

        # Create subheader style
        subheader_style = styles.add_style('CustomSubHeader', WD_STYLE_TYPE.PARAGRAPH)
        subheader_style.font.size = Pt(12)
        subheader_style.font.bold = True
        subheader_style.font.color.rgb = RGBColor(0, 51, 102)
        
        # Document Header
        print("Adding document header...")
        header = doc.add_paragraph()
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = header.add_run(company_name)
        company_run.bold = True
        company_run.font.size = Pt(20)

        # Dynamic Title based on policy type
        print("Adding dynamic title section...")
        title_para = doc.add_paragraph(style='CustomHeader')
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.add_run('{#policy_type == "auto"}Automobile Insurance Policy{/policy_type == "auto"}')
        title_para.add_run('{#policy_type == "home"}Homeowner\'s Insurance Policy{/policy_type == "home"}')
        title_para.add_run('{#policy_type == "life"}Life Insurance Policy{/policy_type == "life"}')

        # Basic Information Section
        print("Adding policy information table...")
        doc.add_paragraph("Policy Information", style='CustomHeader')
        policy_table = doc.add_table(rows=7, cols=2)
        policy_table.style = 'Table Grid'
        
        policy_fields = [
            ("Policy Number:", "{policy_number}"),
            ("Issue Date:", "{issue_date}"),
            ("Effective Date:", "{effective_date}"),
            ("Expiration Date:", "{expiration_date}"),
            ("Policy Status:", "{status}"),
            ("Annual Premium:", "${premiumDetails.annualPremium}"),
            ("Payment Frequency:", "{premiumDetails.paymentFrequency}")
        ]
        
        for i, (field, value) in enumerate(policy_fields):
            row_cells = policy_table.rows[i].cells
            row_cells[0].text = field
            row_cells[1].text = value

        doc.add_paragraph()  # Spacing

        # Policyholder Information
        print("Adding policyholder section...")
        doc.add_paragraph("Policyholder Information", style='CustomHeader')
        policyholder_table = doc.add_table(rows=6, cols=2)
        policyholder_table.style = 'Table Grid'
        
        policyholder_fields = [
            ("Full Name:", "{full_name}"),
            ("Address:", "{address}"),
            ("City, State, ZIP:", "{city_state_zip}"),
            ("Phone:", "{phone_number}"),
            ("Email:", "{email_address}"),
            ("Date of Birth:", "{date_of_birth}")
        ]
        
        for i, (field, value) in enumerate(policyholder_fields):
            row_cells = policyholder_table.rows[i].cells
            row_cells[0].text = field
            row_cells[1].text = value

        doc.add_paragraph()  # Spacing

        print("Adding policy-specific sections...")
        # Coverage Details Header
        doc.add_paragraph("Coverage Details", style='CustomHeader')

        # Auto Insurance Details
        auto_section = doc.add_paragraph()
        auto_section.add_run('{#policy_type == "auto"}\n')
        auto_section.add_run("Vehicle Information:\n")
        auto_section.add_run("Make: {coverageDetails.vehicleInfo.make}\n")
        auto_section.add_run("Model: {coverageDetails.vehicleInfo.model}\n")
        auto_section.add_run("Year: {coverageDetails.vehicleInfo.year}\n")
        auto_section.add_run("VIN: {coverageDetails.vehicleInfo.vin}\n\n")
        auto_section.add_run("Coverage:\n")
        auto_section.add_run("- Liability Coverage: {coverage_limit}\n")
        auto_section.add_run("- Collision Deductible: {deductible_amount}\n")
        auto_section.add_run("{/policy_type == \"auto\"}")

        # Home Insurance Details
        print("Adding home insurance details...")
        home_section = doc.add_paragraph()
        home_section.add_run('{#policy_type == "home"}\n')
        home_section.add_run("Property Information:\n")
        home_section.add_run("Construction Year: {coverageDetails.propertyInfo.constructionYear}\n")
        home_section.add_run("Square Feet: {coverageDetails.propertyInfo.squareFeet}\n")
        home_section.add_run("Construction Type: {coverageDetails.propertyInfo.constructionType}\n\n")
        home_section.add_run("Coverage:\n")
        home_section.add_run("- Dwelling Coverage: {coverage_limit}\n")
        home_section.add_run("- Personal Property: {coverageDetails.personalPropertyLimit}\n")
        home_section.add_run("{/policy_type == \"home\"}")

        # Life Insurance Details
        life_section = doc.add_paragraph()
        life_section.add_run('{#policy_type == "life"}\n')
        life_section.add_run("Coverage Details:\n")
        life_section.add_run("- Death Benefit: {coverage_limit}\n")
        life_section.add_run("- Coverage Type: {coverage_description}\n")
        life_section.add_run("- Term Length: {coverageDetails.termLength}\n")
        life_section.add_run("{/policy_type == \"life\"}")

        # Premium Information
        print("Adding premium information...")
        doc.add_paragraph("Premium Information", style='CustomSubHeader')
        premium_section = doc.add_paragraph()
        premium_section.add_run("Annual Premium: ${premiumDetails.annualPremium}\n")
        premium_section.add_run("Payment Frequency: {premiumDetails.paymentFrequency}\n")
        premium_section.add_run("Next Payment Due: {premiumDetails.nextPaymentDue}\n")
        premium_section.add_run('{#premiumDetails.discount > 0}')
        premium_section.add_run("Applied Discount: ${premiumDetails.discount}\n")
        premium_section.add_run('{/premiumDetails.discount > 0}')

        # Status-based Messages
        doc.add_paragraph("Policy Status", style='CustomSubHeader')
        status_section = doc.add_paragraph()
        status_section.add_run('{#status == "active"}')
        status_section.add_run("Your policy is currently active and in force.\n")
        status_section.add_run('{/status == "active"}')
        status_section.add_run('{#status == "pending"}')
        status_section.add_run("Your policy is pending activation. Please contact our office.\n")
        status_section.add_run('{/status == "pending"}')
        status_section.add_run('{#status == "expired"}')
        status_section.add_run("Your policy has expired. Please contact us immediately.\n")
        status_section.add_run('{/status == "expired"}')

        # VIP Section
        vip_section = doc.add_paragraph()
        vip_section.add_run('{#coverage_limit_number >= 500000}')
        vip_section.add_run("\nAs a premium policyholder, you have access to our VIP support line: 1-800-VIP-SUPPORT")
        vip_section.add_run('{/coverage_limit_number >= 500000}')

        # Declarations and Signatures
        print("Adding signature section...")
        doc.add_paragraph("\nDeclarations and Signatures", style='CustomHeader')
        signature_table = doc.add_table(rows=4, cols=2)
        signature_table.style = 'Table Grid'
        
        signature_fields = [
            ("Insurance Representative:", "________________________"),
            ("Date:", "{representative_signature_date}"),
            ("Policyholder Signature:", "________________________"),
            ("Date:", "{policyholder_signature_date}")
        ]
        
        for i, (field, value) in enumerate(signature_fields):
            signature_table.rows[i].cells[0].text = field
            signature_table.rows[i].cells[1].text = value

        # Add footer with page numbers
        section = doc.sections[0]
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.text = "Page "
        footer_para.add_run()

        print("Template creation completed successfully!")
        return doc

    except Exception as e:
        print(f"Error creating template: {str(e)}")
        raise

def save_template(company_name="Sample Insurance Co.", output_path=None):
    """Creates and saves an insurance document template."""
    try:
        # Create default output path if none provided
        if output_path is None:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = os.path.join("output", f"insurance_template_{timestamp}.docx")

        # Ensure output directory exists
        ensure_output_directory(os.path.dirname(output_path))

        # Create and save the template
        doc = create_insurance_template(company_name)
        doc.save(output_path)
        
        print(f"\nTemplate saved successfully to: {output_path}")
        return output_path

    except Exception as e:
        print(f"\nError saving template: {str(e)}")
        raise

if __name__ == "__main__":
    try:
        print("Starting Insurance Template Generator...")
        
        # You can customize these values
        COMPANY_NAME = "Your Insurance Company Name"
        OUTPUT_PATH = "output/insurance_template.docx"
        
        template_path = save_template(
            company_name=COMPANY_NAME,
            output_path=OUTPUT_PATH
        )
        
        print("\nTemplate generation completed successfully!")
        print(f"You can find your template at: {template_path}")
        
    except Exception as e:
        print(f"\nTemplate generation failed: {str(e)}")
        print("\nPlease ensure you have the required dependencies installed:")
        print("pip install python-docx")