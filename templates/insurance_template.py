from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_insurance_quote_template():
    """
    Creates a detailed, fillable insurance quote template with clear placeholders.
    This template uses curly braces {} to indicate where information needs to be filled in,
    making it easier for users to identify and replace placeholder text.
    """
    doc = Document()
    
    # First, let's set up our custom styles for a professional appearance
    styles = {
        'CustomHeader': {
            'size': 16,
            'color': RGBColor(0, 51, 102),  # Professional dark blue
            'bold': True
        },
        'CustomSubHeader': {
            'size': 14,
            'color': RGBColor(51, 51, 51),  # Dark gray
            'bold': True
        },
        'CustomPrompt': {
            'size': 11,
            'color': RGBColor(128, 128, 128),  # Medium gray
            'italic': True
        },
        'CustomNormal': {
            'size': 11,
            'color': RGBColor(0, 0, 0),  # Black
            'bold': False
        }
    }
    
    # Apply our custom styles to the document
    for style_name, properties in styles.items():
        try:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.font.size = Pt(properties['size'])
            style.font.bold = properties.get('bold', False)
            style.font.italic = properties.get('italic', False)
            style.font.color.rgb = properties['color']
        except ValueError:
            # If style already exists, get it and modify it
            style = doc.styles[style_name]
            style.font.size = Pt(properties['size'])
            style.font.bold = properties.get('bold', False)
            style.font.italic = properties.get('italic', False)
            style.font.color.rgb = properties['color']

    # Company Header Section
    doc.add_paragraph("INSURANCE QUOTE", style='CustomHeader').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("{Company Logo}", style='CustomPrompt').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("{Company Name}", style='CustomPrompt').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("{Company License Number}", style='CustomPrompt').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Quote Information Section
    doc.add_paragraph("QUOTE INFORMATION", style='CustomSubHeader')
    quote_info = doc.add_paragraph(style='CustomNormal')
    quote_info.add_run("Quote Reference: ").bold = True
    quote_info.add_run("{Quote Reference Number}")
    quote_info.add_run("\nDate Generated: ").bold = True
    quote_info.add_run(datetime.now().strftime('%Y-%m-%d'))
    quote_info.add_run("\nValid Until: ").bold = True
    quote_info.add_run("{Validity Date}")
    
    # Agent Information
    doc.add_paragraph("AGENT INFORMATION", style='CustomSubHeader')
    agent_info = doc.add_paragraph(style='CustomNormal')
    agent_info.add_run("Name: ").bold = True
    agent_info.add_run("{Agent Full Name}")
    agent_info.add_run("\nLicense Number: ").bold = True
    agent_info.add_run("{Agent License Number}")
    agent_info.add_run("\nContact: ").bold = True
    agent_info.add_run("{Agent Phone and Email}")
    
    # Client Information Section
    doc.add_paragraph("CLIENT INFORMATION", style='CustomSubHeader')
    doc.add_paragraph("Personal Details", style='CustomNormal').bold = True
    
    # Create a professional table for client information
    client_table = doc.add_table(rows=8, cols=2)
    client_table.style = 'Table Grid'
    
    # Define client information fields with curly brace placeholders
    client_fields = [
        ("Full Name:", "{Client's Full Legal Name}"),
        ("Date of Birth:", "{MM/DD/YYYY}"),
        ("Address:", "{Complete Mailing Address}"),
        ("Phone Number:", "{Primary Contact Number}"),
        ("Email:", "{Email Address}"),
        ("Occupation:", "{Current Occupation}"),
        ("Current Insurance:", "{Current Provider if any}"),
        ("Policy Type:", "{Requested Policy Type}")
    ]
    
    # Fill the client information table
    for i, (field, placeholder) in enumerate(client_fields):
        row = client_table.rows[i]
        row.cells[0].text = field
        row.cells[1].text = placeholder
    
    doc.add_paragraph()  # Add spacing
    
    # Coverage Details Section
    doc.add_paragraph("COVERAGE DETAILS", style='CustomSubHeader')
    coverage_table = doc.add_table(rows=1, cols=4)
    coverage_table.style = 'Table Grid'
    
    # Set header row
    header_cells = coverage_table.rows[0].cells
    for cell, header in zip(header_cells, ['Coverage Type', 'Amount', 'Deductible', 'Premium']):
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    # Add sample rows with curly brace placeholders
    for _ in range(5):
        row_cells = coverage_table.add_row().cells
        row_cells[0].text = "{Coverage Type}"
        row_cells[1].text = "{Amount}"
        row_cells[2].text = "{Deductible}"
        row_cells[3].text = "{Premium}"
    
    # Premium Summary Section
    doc.add_paragraph("PREMIUM SUMMARY", style='CustomSubHeader')
    premium_table = doc.add_table(rows=4, cols=2)
    premium_table.style = 'Table Grid'
    
    premium_items = [
        ("Base Premium:", "{Base Premium Amount}"),
        ("Discounts:", "{List of Applied Discounts}"),
        ("Additional Fees:", "{Additional Fees}"),
        ("Total Annual Premium:", "{Total Premium Amount}")
    ]
    
    for i, (item, value) in enumerate(premium_items):
        cells = premium_table.rows[i].cells
        cells[0].text = item
        cells[1].text = value
    
    # Payment Options Section
    doc.add_paragraph("PAYMENT OPTIONS", style='CustomSubHeader')
    payment_table = doc.add_table(rows=3, cols=2)
    payment_table.style = 'Table Grid'
    
    payment_options = [
        ("Annual:", "{Annual Payment Details}"),
        ("Semi-Annual:", "{Semi-Annual Payment Details}"),
        ("Monthly:", "{Monthly Payment Details}")
    ]
    
    for i, (option, details) in enumerate(payment_options):
        cells = payment_table.rows[i].cells
        cells[0].text = option
        cells[1].text = details
    
    # Terms and Conditions
    doc.add_paragraph("TERMS AND CONDITIONS", style='CustomSubHeader')
    doc.add_paragraph("{Standard Terms and Conditions}", style='CustomPrompt')
    
    # Signature Section
    doc.add_paragraph("AUTHORIZATION", style='CustomSubHeader')
    
    signatures = [
        "_________________________    ______________",
        "Client Signature             Date",
        "",
        "_________________________    ______________",
        "Agent Signature             Date"
    ]
    
    for line in signatures:
        doc.add_paragraph(line, style='CustomNormal')
    
    # Footer with placeholders
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "{Company Name} | {Phone} | {Email} | {Website}"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the template with a descriptive name
    template_name = 'Insurance_Quote_Template.docx'
    doc.save(template_name)
    
    return f"Template created successfully as '{template_name}'"

if __name__ == "__main__":
    try:
        result = create_insurance_quote_template()
        print(result)
    except Exception as e:
        print(f"An error occurred: {str(e)}")