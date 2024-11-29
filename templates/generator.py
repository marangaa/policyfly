from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime

def create_insurance_template(company_name="Sample Insurance Co."):
    """
    Creates an insurance document template with proper template tag handling for Docxtemplater.
    Uses {tag} syntax for template variables that will be replaced with actual values.
    """
    doc = Document()
    
    # Set up default styles
    styles = doc.styles
    
    # Create custom style for headers
    header_style = styles.add_style('CustomHeader', WD_STYLE_TYPE.PARAGRAPH)
    header_style.font.size = Pt(16)
    header_style.font.bold = True
    header_style.font.color.rgb = RGBColor(0, 51, 102)
    
    # Create custom style for subheaders
    subheader_style = styles.add_style('CustomSubHeader', WD_STYLE_TYPE.PARAGRAPH)
    subheader_style.font.size = Pt(12)
    subheader_style.font.bold = True
    
    # Add company header as a normal text (not a template variable)
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_run = header.add_run(company_name)
    company_run.bold = True
    company_run.font.size = Pt(20)
    
    # Add document title
    title = doc.add_paragraph("Insurance Policy Document", style='CustomHeader')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add policy information section
    doc.add_paragraph("Policy Information", style='CustomHeader')
    policy_table = doc.add_table(rows=5, cols=2)
    policy_table.style = 'Table Grid'
    
    # Define policy information fields with properly escaped template tags
    policy_fields = [
        ("Policy Number:", "{policy_number}"),
        ("Date Issued:", "{issue_date}"),
        ("Effective Date:", "{effective_date}"),
        ("Expiration Date:", "{expiration_date}"),
        ("Policy Type:", "{policy_type}")
    ]
    
    for i, (field, value) in enumerate(policy_fields):
        row_cells = policy_table.rows[i].cells
        row_cells[0].text = field
        row_cells[1].text = value
    
    doc.add_paragraph()  # Add spacing
    
    # Policyholder Information
    doc.add_paragraph("Policyholder Information", style='CustomHeader')
    policyholder_table = doc.add_table(rows=6, cols=2)
    policyholder_table.style = 'Table Grid'
    
    # Define policyholder fields
    policyholder_fields = [
        ("Full Name:", "{full_name}"),
        ("Address:", "{address}"),
        ("City, State, ZIP:", "{city_state_zip}"),
        ("Phone:", "{phone_number}"),
        ("Email:", "{email_address}"),
        ("Date of Birth:", "{date_of_birth}")
    ]
    
    for i, (field, value) in enumerate(policyholder_fields):
        row_cells = policyholder_table.rows[i].cells
        row_cells[0].text = field
        row_cells[1].text = value
    
    doc.add_paragraph()  # Add spacing
    
    # Coverage Details
    doc.add_paragraph("Coverage Details", style='CustomHeader')
    coverage_table = doc.add_table(rows=1, cols=3)
    coverage_table.style = 'Table Grid'
    
    # Set up coverage table headers
    headers = ["Coverage Type", "Limit", "Deductible"]
    for i, header in enumerate(headers):
        cell = coverage_table.cell(0, i)
        cell.text = header
        cell.paragraphs[0].style = 'CustomSubHeader'
    
    # Add sample coverage row
    row = coverage_table.add_row()
    row.cells[0].text = "{coverage_description}"
    row.cells[1].text = "{coverage_limit}"
    row.cells[2].text = "{deductible_amount}"
    
    doc.add_paragraph()  # Add spacing
    
    # Terms and Conditions
    doc.add_paragraph("Terms and Conditions", style='CustomHeader')
    terms = doc.add_paragraph()
    terms.add_run("{terms_and_conditions}")
    
    # Declarations and Signatures
    doc.add_paragraph("Declarations and Signatures", style='CustomHeader')
    signature_table = doc.add_table(rows=4, cols=2)
    signature_table.style = 'Table Grid'
    
    # Define signature fields
    signature_fields = [
        ("Insurance Representative:", "________________________"),
        ("Date:", "{representative_signature_date}"),
        ("Policyholder Signature:", "________________________"),
        ("Date:", "{policyholder_signature_date}")
    ]
    
    for i, (field, value) in enumerate(signature_fields):
        row_cells = signature_table.rows[i].cells
        row_cells[0].text = field
        row_cells[1].text = value
    
    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.text = "Page "
    footer_para.add_run()  # Word will automatically replace this with the page number
    
    return doc

def save_template(company_name="Sample Insurance Co.", output_path="insurance_template.docx"):
    """
    Creates and saves an insurance document template.
    
    Args:
        company_name (str): Name of the insurance company
        output_path (str): Path where the document should be saved
    """
    doc = create_insurance_template(company_name)
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    # Example usage
    template_path = save_template(
        company_name="Acme Insurance Corporation",
        output_path="insurance_policy_template.docx"
    )
    print(f"Template saved to: {template_path}")