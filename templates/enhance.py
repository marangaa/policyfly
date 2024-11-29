from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import locale

# Set locale for currency formatting
locale.setlocale(locale.LC_ALL, 'en_US.UTF-8')

class InsuranceQuoteTemplate:
    """
    A comprehensive template generator for insurance quotes.
    
    This class provides methods to create professional insurance quotes with
    customizable sections, formatting, and content. It supports multiple
    insurance types and can be extended for specific company needs.
    """
    
    def __init__(self, company_info=None):
        """Initialize the quote template with company information."""
        self.doc = Document()
        self.company_info = company_info or {}
        self._setup_styles()
        self._setup_page_format()
    
    def _setup_styles(self):
        """Configure document styles for consistent formatting."""
        # Header style for main titles
        header_style = self.doc.styles.add_style('Header Style', WD_STYLE_TYPE.PARAGRAPH)
        header_style.font.size = Pt(16)
        header_style.font.bold = True
        header_style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue for professionalism
        
        # Subheader style for section titles
        subheader_style = self.doc.styles.add_style('Subheader Style', WD_STYLE_TYPE.PARAGRAPH)
        subheader_style.font.size = Pt(14)
        subheader_style.font.bold = True
        subheader_style.font.color.rgb = RGBColor(51, 51, 51)  # Dark gray
        
        # Normal text style
        normal_style = self.doc.styles.add_style('Normal Style', WD_STYLE_TYPE.PARAGRAPH)
        normal_style.font.size = Pt(11)
        normal_style.font.name = 'Arial'
        
        # Table style for coverage details
        table_style = self.doc.styles.add_style('Table Style', WD_STYLE_TYPE.TABLE)
        table_style.font.size = Pt(10)
        table_style.font.name = 'Arial'
    
    def _setup_page_format(self):
        """Configure page formatting including margins."""
        section = self.doc.sections[0]
        section.page_margin_left = Inches(1)
        section.page_margin_right = Inches(1)
        section.page_margin_top = Inches(1)
        section.page_margin_bottom = Inches(1)
    
    def _format_currency(self, amount):
        """Format number as currency string with proper handling of invalid inputs."""
        try:
            return locale.currency(float(amount), grouping=True)
        except (ValueError, TypeError):
            return str(amount)
    
    def _get_policy_notes(self, policy_type):
        """Return policy-specific notes based on insurance type."""
        notes = {
            'Auto': 'Coverage applies to listed vehicles only. Additional drivers must be registered.',
            'Home': 'Coverage based on property evaluation and risk assessment.',
            'Life': 'Coverage subject to medical examination and history review.',
            'Business': 'Coverage applies to declared business activities and locations.'
        }
        return notes.get(policy_type, 'Standard terms and conditions apply.')
    
    def _add_page_number(self, paragraph):
        """Add page numbers to the document footer."""
        page_num_run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        page_num_run._r.append(fldChar1)
        page_num_run._r.append(instrText)
        page_num_run._r.append(fldChar2)
    
    def add_company_header(self, logo_path=None):
        """Add company header with logo and registration information."""
        if logo_path:
            self.doc.add_picture(logo_path, width=Inches(2))
        
        company_header = self.doc.add_paragraph(style='Header Style')
        company_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_header.add_run(self.company_info.get('name', 'Insurance Company')).bold = True
        
        if 'license_number' in self.company_info:
            reg_info = self.doc.add_paragraph(style='Normal Style')
            reg_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
            reg_info.add_run(f"Licensed Insurance Provider - {self.company_info['license_number']}")
        
        self.doc.add_paragraph()
    
    def add_quote_info(self, quote_data):
        """Add detailed quote reference information."""
        quote_info = self.doc.add_paragraph(style='Normal Style')
        
        quote_info.add_run('Quote Reference: ').bold = True
        quote_info.add_run(quote_data.get('reference', 'TBD'))
        quote_info.add_run('\nDate Generated: ').bold = True
        quote_info.add_run(datetime.now().strftime('%Y-%m-%d'))
        quote_info.add_run('\nValid Until: ').bold = True
        quote_info.add_run(quote_data.get('valid_until', 'N/A'))
        
        if 'agent' in quote_data:
            quote_info.add_run('\n\nInsurance Agent: ').bold = True
            quote_info.add_run(quote_data['agent'].get('name', ''))
            quote_info.add_run('\nAgent License: ').bold = True
            quote_info.add_run(quote_data['agent'].get('license', ''))
            quote_info.add_run('\nContact: ').bold = True
            quote_info.add_run(quote_data['agent'].get('contact', ''))
        
        self.doc.add_paragraph()
    
    def add_client_info(self, client_data):
        """Add comprehensive client information section."""
        self.doc.add_paragraph('Client Information', style='Subheader Style')
        
        table = self.doc.add_table(rows=1, cols=2)
        left_cell, right_cell = table.rows[0].cells
        
        personal_info = left_cell.add_paragraph(style='Normal Style')
        personal_info.add_run('Personal Details\n').bold = True
        for key in ['Name', 'Date of Birth', 'Address', 'Phone', 'Email']:
            personal_info.add_run(f'{key}: ').bold = True
            personal_info.add_run(f"{client_data.get(key, 'N/A')}\n")
        
        policy_info = right_cell.add_paragraph(style='Normal Style')
        policy_info.add_run('Policy Information\n').bold = True
        for key in ['Policy Type', 'Current Provider', 'Claims History', 'Risk Level']:
            policy_info.add_run(f'{key}: ').bold = True
            policy_info.add_run(f"{client_data.get(key, 'N/A')}\n")
        
        self.doc.add_paragraph()
    
    def add_coverage_details(self, coverage_items, policy_type):
        """Add detailed coverage information with customization based on policy type."""
        self.doc.add_paragraph(f'{policy_type} Coverage Details', style='Subheader Style')
        
        notes = self.doc.add_paragraph(style='Normal Style')
        notes.add_run('Coverage Overview: ').bold = True
        notes.add_run(self._get_policy_notes(policy_type))
        
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        headers = ['Coverage Type', 'Coverage Amount', 'Deductible', 'Annual Premium']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for item in coverage_items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('type', '')
            row_cells[1].text = self._format_currency(item.get('amount', 0))
            row_cells[2].text = self._format_currency(item.get('deductible', 0))
            row_cells[3].text = self._format_currency(item.get('premium', 0))
        
        self.doc.add_paragraph()
    
    def add_terms_and_conditions(self, terms, disclaimers=None):
        """Add terms, conditions, and disclaimers."""
        self.doc.add_paragraph('Terms and Conditions', style='Subheader Style')
        
        terms_para = self.doc.add_paragraph(style='Normal Style')
        terms_para.add_run('By accepting this quote, you agree to the following terms:\n\n')
        for term in terms:
            terms_para.add_run(f'• {term}\n')
        
        if disclaimers:
            self.doc.add_paragraph('Important Disclaimers', style='Subheader Style')
            disclaimer_para = self.doc.add_paragraph(style='Normal Style')
            disclaimer_para.add_run('Please note:\n\n')
            for disclaimer in disclaimers:
                disclaimer_para.add_run(f'* {disclaimer}\n')
        
        self.doc.add_paragraph()
    
    def add_premium_summary(self, premium_data):
        """Add comprehensive premium summary with payment options."""
        self.doc.add_paragraph('Premium Summary', style='Subheader Style')
        
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        left_cell = table.rows[0].cells[0]
        left_cell.add_paragraph('Premium Breakdown', style='Normal Style').bold = True
        
        breakdown = premium_data.get('breakdown', {})
        for item, amount in breakdown.items():
            para = left_cell.add_paragraph(style='Normal Style')
            para.add_run(f'{item}: ').bold = True
            para.add_run(self._format_currency(amount))
        
        right_cell = table.rows[0].cells[1]
        right_cell.add_paragraph('Payment Options', style='Normal Style').bold = True
        
        options = premium_data.get('payment_options', [])
        for option in options:
            para = right_cell.add_paragraph(style='Normal Style')
            para.add_run(f"{option['term']}: ").bold = True
            para.add_run(f"{option['description']}")
        
        self.doc.add_paragraph()
    
    def add_footer(self, include_page_numbers=True):
        """Add professional footer with company contact information and optional page numbers."""
        footer = self.doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        
        footer_para.text = (
            f"{self.company_info.get('name', 'Insurance Company')} | "
            f"Tel: {self.company_info.get('phone', '')} | "
            f"Email: {self.company_info.get('email', '')}"
        )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if include_page_numbers:
            self._add_page_number(footer_para)
    
    def save_document(self, filename):
        """Save the generated quote document."""
        self.doc.save(filename)


def create_sample_quote():
    """Create a sample insurance quote with all available features."""
    # Company information
    company_info = {
        'name': 'ABC Insurance Company',
        'license_number': 'INS-2024-123456',
        'phone': '(800) 555-0123',
        'email': 'quotes@abcinsurance.com'
    }
    
    template = InsuranceQuoteTemplate(company_info)
    
    # Quote information
    quote_data = {
        'reference': 'QT-2024-001',
        'valid_until': '2024-12-31',
        'agent': {
            'name': 'Jane Smith',
            'license': 'AG123456',
            'contact': '(555) 555-5555'
        }
    }
    
    # Client information
    client_data = {
        'Name': 'John Doe',
        'Date of Birth': '1980-01-15',
        'Address': '123 Main Street, Anytown, ST 12345',
        'Phone': '(555) 555-5555',
        'Email': 'john.doe@email.com',
        'Policy Type': 'Auto Insurance',
        'Current Provider': 'Previous Insurance Co',
        'Claims History': 'No claims in past 5 years',
        'Risk Level': 'Low'
    }
    
    # Coverage details for auto insurance
    coverage_items = [
        {
            'type': 'Liability - Bodily Injury',
            'amount': 300000,
            'deductible': 0,
            'premium': 800
        },
        {
            'type': 'Liability - Property Damage',
            'amount': 100000,
            'deductible': 0,
            'premium': 400
        },
        {
            'type': 'Collision Coverage',
            'amount': 50000,
            'deductible': 500,
            'premium': 600
        },
        {
            'type': 'Comprehensive Coverage',
            'amount': 50000,
            'deductible': 250,
            'premium': 300
        },
        {
            'type': 'Personal Injury Protection',
            'amount': 25000,
            'deductible': 0,
            'premium': 200
        }
    ]
    
    # Premium summary data
    premium_data = {
        'breakdown': {
            'Base Premium': 2300,
            'Safe Driver Discount': -230,
            'Multi-Policy Discount': -115,
            'Policy Fees': 45
        },
        'payment_options': [
            {
                'term': 'Annual',
                'description': 'Single payment of $2,000 (Save 10%)'
            },
            {
                'term': 'Semi-Annual',
                'description': 'Two payments of $1,050'
            },
            {
                'term': 'Monthly',
                'description': 'Twelve payments of $183.33'
            }
        ]
    }
    
    # Terms and conditions
    terms = [
        'Coverage begins upon receipt of first payment',
        'This quote is based on the information provided and subject to verification',
        '30-day notice required for policy cancellation',
        'Claims must be reported within 24 hours of incident',
        'Coverage is subject to policy terms, conditions, and exclusions',
        'Deductibles apply per incident as specified in coverage details',
        'Premium rates are subject to change based on underwriting review',
        'All coverage limits are on a per-occurrence basis unless otherwise specified'
    ]
    
    # Important disclaimers
    disclaimers = [
        'This quote is not a binding contract and is subject to underwriting review',
        'Rates may change based on final verification of provided information',
        'Additional fees may apply based on state regulations and payment method',
        'Coverage exclusions may apply. Please refer to policy documents for complete details',
        'This quote assumes all provided information is accurate and complete'
    ]
    
    # Generate the complete quote document
    template.add_company_header()  # Add logo_path parameter if you have a company logo
    template.add_quote_info(quote_data)
    template.add_client_info(client_data)
    template.add_coverage_details(coverage_items, 'Auto')
    template.add_terms_and_conditions(terms, disclaimers)
    template.add_premium_summary(premium_data)
    template.add_footer(include_page_numbers=True)
    
    # Save the final document
    template.save_document('enhanced_insurance_quote.docx')


if __name__ == "__main__":
    try:
        create_sample_quote()
        print("Insurance quote document generated successfully!")
    except Exception as e:
        print(f"An error occurred while generating the quote: {str(e)}")